import os
import re
import requests
import bs4
import docx

link_list = "links.txt"
folder_name = "cifras"
document_model = "modelo.docx"


with open(link_list, "r") as links:
    content = links.read().split("\n")
    for link in content:
        if link:
            mobile_url = re.sub(r"(www)", "m", link)
            r = requests.get(mobile_url)
            r.raise_for_status()
            page_html = bs4.BeautifulSoup(r.text, "html.parser")
            page_title = re.search(r"(.*) - (.*) - (.*)", page_html.title.string)
            chords_title = page_title.group(2) + " - " + page_title.group(1)
            chords_body = str(page_html.pre)
            chords_document = docx.Document(document_model)
            document_header = chords_document.sections[0].header
            document_header.paragraphs[0].text = chords_title
            p = chords_document.add_paragraph()
            for a in page_html.find_all("span", "_1sHaH"):
                p.add_run(a.get_text() + "\n")
            p.add_run("\n")
            for i in chords_body.split("\n"):
                if re.search(r"<b*?>", i):
                    note = re.sub(r"<.*?>", "", i)
                    p.add_run(note + "\n").bold = True
                else:
                    lyric = re.sub(r"<.*?>", "", i)
                    p.add_run(lyric + "\n")
            p.add_run("\n[Acordes]\n\n")
            for c in page_html.find_all("div", "chord"):
                p.add_run(c.get("data-mount") + " - " + c.find("strong").text + "\n")
            os.makedirs(folder_name, exist_ok=True)
            file_name = re.sub(r"\W+", "_", chords_title)
            chords_document.save(folder_name + "/" + file_name + ".docx")
            print("Arquivo criado: " + chords_title)
print("Operação concluída!")
